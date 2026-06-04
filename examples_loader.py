"""
Загружает примеры документов из папок examples/.
"""
import os
import logging

logger = logging.getLogger(__name__)

BASE_DIR = os.path.join(os.path.dirname(__file__), "examples")

FOLDER_MAP = {
    "cleaning": "cleaning",
    "it": "it",
    "repair": "repair",
    "criteria": "criteria",
}


def load_examples(category: str) -> list[str]:
    folder = FOLDER_MAP.get(category, category)
    path = os.path.join(BASE_DIR, folder)
    if not os.path.exists(path):
        logger.warning(f"Папка не найдена: {path}")
        return []

    texts = []
    for filename in sorted(os.listdir(path)):
        if not filename.lower().endswith(".docx"):
            continue
        filepath = os.path.join(path, filename)
        text = _read_docx(filepath)
        if text:
            texts.append(text)
            logger.info(f"Загружен: {filename} ({len(text)} симв.)")

    logger.info(f"Категория '{category}': загружено {len(texts)} файлов")
    return texts


def _read_docx(filepath: str) -> str:
    try:
        from docx import Document
        doc = Document(filepath)
        parts = []
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)
        # Таблицы тоже читаем
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as e:
        logger.error(f"Ошибка чтения {filepath}: {e}")
        return ""
