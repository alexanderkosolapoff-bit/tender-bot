"""
Генерация Word и PDF документов.
Критерии допуска — таблица: № п/п | Наименование критерия | Значение | Подтверждающий документ
Word и PDF строятся из одних и тех же данных — одинаковый результат.
"""
import os, re, tempfile, logging
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

# Шрифты — ищем в папке fonts/, потом в корне, потом системные
_BASE = os.path.dirname(os.path.abspath(__file__))
def _find_font(name):
    for p in [
        os.path.join(_BASE, "fonts", name),
        os.path.join(_BASE, name),
        f"/usr/share/fonts/truetype/dejavu/{name}",
        f"/usr/share/fonts/truetype/liberation/{name.replace('DejaVuSerif', 'LiberationSerif').replace('-Bold','-Bold').replace('-Italic','-Italic')}",
    ]:
        if os.path.exists(p): return p
    return None

FONT_REG  = _find_font("DejaVuSerif.ttf")
FONT_BOLD = _find_font("DejaVuSerif-Bold.ttf")
FONT_ITAL = _find_font("DejaVuSerif-Italic.ttf")

# Стандартный текст после таблицы критериев
CRITERIA_FOOTER = [
    "Для подтверждения соответствия критериям допуска участник обязан предоставить документы, "
    "указанные в графе «Подтверждающий документ», в составе заявки на участие в закупке.",
    "Все представленные документы должны быть действительны на дату подачи заявки.",
    "Организатор закупки вправе запросить дополнительные документы для подтверждения "
    "соответствия участника установленным критериям.",
    "Примечание: участник, не соответствующий хотя бы одному из критериев допуска, "
    "не допускается к дальнейшему рассмотрению заявки.",
]


def _strip_md(t):
    t = re.sub(r'\*\*(.+?)\*\*', r'\1', t)
    t = re.sub(r'\*(.+?)\*',     r'\1', t)
    t = re.sub(r'`(.+?)`',       r'\1', t)
    t = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', t)
    t = re.sub(r'^#+\s*', '', t)
    return t.strip()


def _is_criteria(content: str) -> bool:
    return bool(re.search(r'КРИТЕРИЙ:', content)) or bool(re.search(r'КРИТЕРИИ ДОПУСКА', content, re.I))


def parse_criteria(text: str) -> list[dict]:
    """Парсит текст в формате КРИТЕРИЙ/ТРЕБОВАНИЕ/ДОКУМЕНТ.
    Устойчив к markdown (**), нумерации в начале, разному регистру."""
    rows = []
    current = {}
    num = 1

    def flush():
        nonlocal current, num
        if current.get("criterion"):
            rows.append({
                "num": str(num),
                "criterion": current.get("criterion", ""),
                "requirement": current.get("requirement", current.get("criterion", "")),
                "document": current.get("document", "По запросу организатора"),
            })
            num += 1
            current = {}

    def clean(s):
        # Убираем markdown ** и * по краям
        s = s.strip()
        s = re.sub(r"^\*+|\*+$", "", s).strip()
        return s

    for raw_line in text.strip().split("\n"):
        line = raw_line.strip()
        if not line:
            flush()
            continue

        # Убираем нумерацию в начале: "1.", "1)", "1.1."
        line_clean = re.sub(r"^\d+[.)]+\s*", "", line).strip()
        # Убираем markdown ** в начале
        line_no_md = re.sub(r"^\*+", "", line_clean).strip()

        up = line_no_md.upper()

        # Ищем ключевое слово в первых 30 символах строки
        head = up[:40]

        if "КРИТЕРИЙ" in head and ":" in line_no_md:
            flush()
            val = line_no_md.split(":", 1)[1]
            current = {"criterion": clean(val)}
        elif "ТРЕБОВАНИЕ" in head and ":" in line_no_md:
            val = line_no_md.split(":", 1)[1]
            current["requirement"] = clean(val)
        elif ("ДОКУМЕНТ" in head or "ПОДТВЕРЖД" in head) and ":" in line_no_md:
            val = line_no_md.split(":", 1)[1]
            current["document"] = clean(val)
        else:
            # Продолжение текущего поля — приписываем
            if current:
                last_key = None
                if "document" in current: last_key = "document"
                elif "requirement" in current: last_key = "requirement"
                elif "criterion" in current: last_key = "criterion"
                if last_key:
                    current[last_key] += " " + clean(line)

    flush()
    return rows


def _is_table_row(s):
    """Строка markdown-таблицы: | a | b | c |"""
    return s.startswith("|") and s.count("|") >= 2

def _is_table_sep(s):
    """Разделитель шапки: |---|---|---|"""
    return bool(re.match(r'^\|[\s:|-]+\|?$', s)) and "-" in s

def _parse_table_row(s):
    """Разбивает | a | b | c | на [a, b, c]"""
    parts = s.strip().strip("|").split("|")
    return [_strip_md(p.strip()) for p in parts]

def parse_content(text: str) -> list[dict]:
    """Парсит текст (ТЗ, сценарий) в элементы, включая markdown-таблицы."""
    els = []
    lines = text.strip().split("\n")
    i = 0
    while i < len(lines):
        s = lines[i].strip()

        # Markdown таблица: несколько строк подряд начинающихся с |
        if _is_table_row(s):
            table_lines = []
            while i < len(lines) and _is_table_row(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1
            # Парсим: первая строка = шапка, вторая = разделитель (пропускаем), остальное = данные
            if len(table_lines) >= 2:
                header = _parse_table_row(table_lines[0])
                data_rows = []
                for tl in table_lines[1:]:
                    if _is_table_sep(tl):
                        continue
                    data_rows.append(_parse_table_row(tl))
                els.append({"type": "table", "header": header, "rows": data_rows})
                continue
            else:
                # Одна строка — не таблица, обычный текст
                els.append({"type": "body", "text": _strip_md(table_lines[0].strip("|"))})
                continue

        if not s:                             els.append({"type": "space"}); i += 1; continue
        if re.match(r'^[-*]{3,}$', s):        els.append({"type": "divider"}); i += 1; continue
        if s.startswith("### "):              els.append({"type": "h3", "text": _strip_md(s[4:])}); i += 1; continue
        if s.startswith("## "):               els.append({"type": "h2", "text": _strip_md(s[3:])}); i += 1; continue
        if s.startswith("# "):                els.append({"type": "h1", "text": _strip_md(s[2:])}); i += 1; continue
        if s.startswith("> "):                els.append({"type": "quote", "text": _strip_md(s[2:])}); i += 1; continue
        if re.match(r'^[-*•]\s', s):          els.append({"type": "bullet", "text": _strip_md(s[2:])}); i += 1; continue
        m = re.match(r'^(\d+)\.\s+(.+)', s)
        if m and len(s) < 200:                els.append({"type": "numbered", "num": m.group(1), "text": _strip_md(m.group(2))}); i += 1; continue
        if s.startswith("**") and (":**" in s or s.endswith("**")):
                                              els.append({"type": "h3", "text": s.strip("*").rstrip(":")}); i += 1; continue
        if re.match(r'^(ТЕХНИЧЕСКОЕ ЗАДАНИЕ|СЦЕНАРИЙ|ОТВЕТН)', s, re.I):
                                              els.append({"type": "title", "text": _strip_md(s)}); i += 1; continue
        els.append({"type": "body", "text": _strip_md(s)}); i += 1
    return els


# ═══════════════════════════════════════════════════════════
# WORD
# ═══════════════════════════════════════════════════════════

def _make_doc():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2); s.bottom_margin = Cm(2)
        s.left_margin = Cm(3); s.right_margin = Cm(1.5)
    n = doc.styles["Normal"]
    n.font.name = "Times New Roman"; n.font.size = Pt(12)
    n.paragraph_format.space_before = Pt(0)
    n.paragraph_format.space_after  = Pt(0)
    n.paragraph_format.line_spacing = Pt(14)
    return doc

def _r(p, text, bold=False, italic=False, size=12, color=None):
    r = p.add_run(text)
    r.font.name = "Times New Roman"; r.font.size = Pt(size)
    r.font.bold = bold; r.font.italic = italic
    if color: r.font.color.rgb = RGBColor(*color)
    return r

def _r_highlighted(p, text, bold=False, italic=False, size=12, color=None):
    """Как _r, но распознаёт маркер ==текст== и подсвечивает эту часть жёлтым.
    Если маркеров нет — ведёт себя идентично _r (один run, без изменений)."""
    parts = re.split(r'(==.+?==)', text)
    last_run = None
    for part in parts:
        if not part:
            continue
        if part.startswith("==") and part.endswith("==") and len(part) > 4:
            run = p.add_run(part[2:-2])
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        else:
            run = p.add_run(part)
        run.font.name = "Times New Roman"; run.font.size = Pt(size)
        run.font.bold = bold; run.font.italic = italic
        if color: run.font.color.rgb = RGBColor(*color)
        last_run = run
    return last_run

def _cell_bg(cell, hex_color):
    tc = cell._tc; pr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pr.append(shd)

def _build_elements_docx(doc, elements):
    for el in elements:
        t = el["type"]
        if t == "space":    doc.add_paragraph()
        elif t == "divider":
            p = doc.add_paragraph(); _r(p, "─" * 50)
        elif t == "title":
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
            _r(p, el["text"], bold=True, size=14)
        elif t == "h1":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
            _r(p, el["text"], bold=True, size=13)
        elif t in ("h2", "numbered"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(3)
            _r_highlighted(p, (el["text"] if t == "h2" else f"{el['num']}. {el['text']}"), bold=True, size=12)
        elif t == "h3":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)
            _r_highlighted(p, el["text"], bold=True, size=12)
        elif t == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(1)
            _r_highlighted(p, el["text"])
        elif t == "quote":
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(1.5)
            _r_highlighted(p, f'"{el["text"]}"', italic=True)
        elif t == "body" and el.get("text"):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1.25)
            _r_highlighted(p, el["text"], size=12)
        elif t == "table":
            _build_generic_table(doc, el["header"], el["rows"])

def _build_generic_table(doc, header, rows):
    """Универсальная таблица для ТЗ — синяя шапка, чередующиеся строки."""
    if not header:
        return
    ncols = len(header)
    table = doc.add_table(rows=1, cols=ncols)
    table.style = "Table Grid"

    # Шапка — синий фон, белый текст
    hdr = table.rows[0]
    for i, h in enumerate(header):
        cell = hdr.cells[i]
        _cell_bg(cell, "1F5C99")
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _r(p, h, bold=True, size=10, color=(255, 255, 255))

    # Данные — чередующиеся строки
    for idx, row in enumerate(rows):
        tr = table.add_row()
        bg = "FFFFFF" if idx % 2 == 0 else "EEF4FB"
        for i in range(ncols):
            cell = tr.cells[i]
            _cell_bg(cell, bg)
            val = row[i] if i < len(row) else ""
            p = cell.paragraphs[0]
            # Первая колонка по центру (обычно номер), остальные слева
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            _r(p, val, size=10)

    doc.add_paragraph()


def _build_criteria_docx(doc, rows):
    """Таблица критериев в Word."""
    headers = ["№ п/п", "Наименование критерия", "Значение", "Подтверждающий документ"]
    widths  = [Cm(1.0), Cm(5.0), Cm(5.2), Cm(5.5)]

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    # Шапка
    hdr = table.rows[0]
    for i, (h, w) in enumerate(zip(headers, widths)):
        cell = hdr.cells[i]; cell.width = w
        _cell_bg(cell, "1F5C99")
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _r(p, h, bold=True, size=10, color=(255, 255, 255))

    # Строки
    for idx, row in enumerate(rows):
        tr = table.add_row()
        bg = "FFFFFF" if idx % 2 == 0 else "EEF4FB"
        vals   = [row["num"], row["criterion"], row["requirement"], row["document"]]
        aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                  WD_ALIGN_PARAGRAPH.LEFT,   WD_ALIGN_PARAGRAPH.LEFT]
        for i, (val, align) in enumerate(zip(vals, aligns)):
            cell = tr.cells[i]; _cell_bg(cell, bg)
            p = cell.paragraphs[0]; p.alignment = align
            _r(p, val, size=10)

    doc.add_paragraph()

    # Стандартный текст после таблицы
    for i, txt in enumerate(CRITERIA_FOOTER):
        p = doc.add_paragraph(txt)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(4)
        italic = (i == len(CRITERIA_FOOTER) - 1)  # Последний абзац — курсив
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(11)
            r.font.italic = italic


async def generate_tz_docx(content, name):
    doc = _make_doc()
    _build_elements_docx(doc, parse_content(content))
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx", prefix="TZ_")
    doc.save(tmp.name); tmp.close()
    return tmp.name

async def generate_criteria_docx(content, name):
    doc = _make_doc()
    # Заголовок
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(10)
    _r(p, "КРИТЕРИИ ДОПУСКА УЧАСТНИКОВ К ЗАКУПКЕ", bold=True, size=14)
    # Таблица
    rows = parse_criteria(content)
    if rows:
        _build_criteria_docx(doc, rows)
    else:
        _build_elements_docx(doc, parse_content(content))
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx", prefix="Crit_")
    doc.save(tmp.name); tmp.close()
    return tmp.name


# ═══════════════════════════════════════════════════════════
# PDF — регистрация шрифтов
# ═══════════════════════════════════════════════════════════

def _reg_fonts():
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfbase.pdfmetrics import registerFontFamily
    if FONT_REG:
        try:
            pdfmetrics.registerFont(TTFont("Cyr",      FONT_REG))
            pdfmetrics.registerFont(TTFont("Cyr-Bold", FONT_BOLD or FONT_REG))
            pdfmetrics.registerFont(TTFont("Cyr-Ital", FONT_ITAL or FONT_REG))
            registerFontFamily("Cyr", normal="Cyr", bold="Cyr-Bold", italic="Cyr-Ital")
            return "Cyr", "Cyr-Bold", "Cyr-Ital"
        except Exception as e:
            logger.warning(f"Font: {e}")
    return "Helvetica", "Helvetica-Bold", "Helvetica-Oblique"


def _make_styles(fn, fn_b, fn_i):
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
    return {
        "title":  ParagraphStyle("TT", fontName=fn_b, fontSize=14, alignment=TA_CENTER, spaceAfter=10, leading=18),
        "h1":     ParagraphStyle("H1", fontName=fn_b, fontSize=13, alignment=TA_LEFT,   spaceAfter=6,  spaceBefore=10, leading=17),
        "h2":     ParagraphStyle("H2", fontName=fn_b, fontSize=12, alignment=TA_LEFT,   spaceAfter=4,  spaceBefore=7,  leading=16),
        "h3":     ParagraphStyle("H3", fontName=fn_b, fontSize=12, alignment=TA_LEFT,   spaceAfter=3,  spaceBefore=5,  leading=16),
        "body":   ParagraphStyle("BD", fontName=fn,   fontSize=12, alignment=TA_JUSTIFY, spaceAfter=4, firstLineIndent=1.25*cm, leading=16),
        "bullet": ParagraphStyle("BL", fontName=fn,   fontSize=12, alignment=TA_LEFT,   spaceAfter=2,  leftIndent=0.8*cm, leading=16),
        "quote":  ParagraphStyle("QT", fontName=fn_i, fontSize=12, alignment=TA_LEFT,   spaceAfter=4,  leftIndent=1.5*cm, leading=16),
        "cell":   ParagraphStyle("CL", fontName=fn,   fontSize=10, alignment=TA_LEFT,   leading=13),
        "cell_c": ParagraphStyle("CC", fontName=fn,   fontSize=10, alignment=TA_CENTER, leading=13),
        "hdr":    ParagraphStyle("CH", fontName=fn_b, fontSize=10, alignment=TA_CENTER, leading=13, textColor=None),
        "foot":   ParagraphStyle("FT", fontName=fn,   fontSize=11, alignment=TA_JUSTIFY, spaceAfter=4, leading=15),
        "note":   ParagraphStyle("NT", fontName=fn_i, fontSize=11, alignment=TA_JUSTIFY, spaceAfter=4, leading=15),
    }


def _build_criteria_pdf(story, rows, st, safe):
    """Таблица критериев в PDF — точно как в Word."""
    from reportlab.platypus import Table, TableStyle, Paragraph
    from reportlab.lib.colors import HexColor
    from reportlab.lib.units import cm

    # Шапка с белым текстом на тёмно-синем фоне
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER
    fn_b = st["hdr"].fontName
    hdr_white = ParagraphStyle("HW", fontName=fn_b, fontSize=10,
                               alignment=TA_CENTER, leading=13,
                               textColor=HexColor("#FFFFFF"))

    headers = ["№ п/п", "Наименование\nкритерия", "Значение", "Подтверждающий\nдокумент"]
    td = [[Paragraph(safe(h), hdr_white) for h in headers]]

    for row in rows:
        td.append([
            Paragraph(safe(row["num"]),         st["cell_c"]),
            Paragraph(safe(row["criterion"]),   st["cell"]),
            Paragraph(safe(row["requirement"]), st["cell"]),
            Paragraph(safe(row["document"]),    st["cell"]),
        ])

    t = Table(td, colWidths=[1.0*cm, 5.0*cm, 5.2*cm, 5.5*cm], repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND",     (0, 0), (-1, 0),  HexColor("#1F5C99")),
        ("GRID",           (0, 0), (-1, -1), 0.5, HexColor("#AAAAAA")),
        ("VALIGN",         (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING",     (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING",  (0, 0), (-1, -1), 4),
        ("LEFTPADDING",    (0, 0), (-1, -1), 3),
        ("RIGHTPADDING",   (0, 0), (-1, -1), 3),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [HexColor("#FFFFFF"), HexColor("#EEF4FB")]),
    ]))
    story.append(t)

    # Стандартный текст — точно как в Word
    from reportlab.platypus import Spacer
    story.append(Spacer(1, 10))
    for i, txt in enumerate(CRITERIA_FOOTER):
        s = st["note"] if i == len(CRITERIA_FOOTER) - 1 else st["foot"]
        story.append(Paragraph(safe(txt), s))


async def generate_pdf(content, name):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.colors import HexColor

    fn, fn_b, fn_i = _reg_fonts()
    st = _make_styles(fn, fn_b, fn_i)

    def safe(t): return (t or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf", prefix="Doc_")
    tmp.close()

    doc = SimpleDocTemplate(tmp.name, pagesize=A4,
        leftMargin=3*cm, rightMargin=1.5*cm, topMargin=2*cm, bottomMargin=2*cm)

    story = []

    if _is_criteria(content):
        # Заголовок
        story.append(Paragraph("КРИТЕРИИ ДОПУСКА УЧАСТНИКОВ К ЗАКУПКЕ", st["title"]))
        rows = parse_criteria(content)
        if rows:
            _build_criteria_pdf(story, rows, st, safe)
        else:
            for line in content.split("\n"):
                if line.strip():
                    story.append(Paragraph(safe(line.strip()), st["body"]))
    else:
        # Обычный документ (ТЗ, сценарий)
        for el in parse_content(content):
            t = el["type"]
            text = safe(el.get("text", ""))
            if t == "space":    story.append(Spacer(1, 6))
            elif t == "divider": story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#888888")))
            elif t == "title":  story.append(Paragraph(text, st["title"]))
            elif t == "h1":     story.append(Paragraph(text, st["h1"]))
            elif t in ("h2", "numbered"):
                story.append(Paragraph((text if t == "h2" else f"{el['num']}. {text}"), st["h2"]))
            elif t == "h3":     story.append(Paragraph(text, st["h3"]))
            elif t == "bullet": story.append(Paragraph(f"• {text}", st["bullet"]))
            elif t == "quote":  story.append(Paragraph(f"«{text}»", st["quote"]))
            elif t == "body" and text.strip():
                story.append(Paragraph(text, st["body"]))
            elif t == "table":
                _build_generic_table_pdf(story, el["header"], el["rows"], st, safe)

    doc.build(story)
    return tmp.name


def _build_generic_table_pdf(story, header, rows, st, safe):
    """Универсальная таблица в PDF."""
    from reportlab.platypus import Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.colors import HexColor
    from reportlab.lib.units import cm
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    if not header:
        return
    ncols = len(header)
    fn_b = st["hdr"].fontName; fn = st["cell"].fontName
    hdr_white = ParagraphStyle("HW2", fontName=fn_b, fontSize=9,
                               alignment=TA_CENTER, leading=11, textColor=HexColor("#FFFFFF"))
    cell_st = ParagraphStyle("CL2", fontName=fn, fontSize=9, alignment=TA_LEFT, leading=11)
    cell_c = ParagraphStyle("CC2", fontName=fn, fontSize=9, alignment=TA_CENTER, leading=11)

    td = [[Paragraph(safe(h), hdr_white) for h in header]]
    for row in rows:
        td.append([Paragraph(safe(row[i] if i < len(row) else ""),
                              cell_c if i == 0 else cell_st) for i in range(ncols)])

    # Ширина колонок — равномерно по доступной ширине (16.5 см)
    total_w = 16.5 * cm
    col_w = [total_w / ncols] * ncols
    # Первая колонка уже если это номер
    if ncols > 1 and all(len(safe(r[0])) <= 5 for r in rows if r):
        col_w[0] = 1.2 * cm
        rest = (total_w - 1.2 * cm) / (ncols - 1)
        col_w = [1.2 * cm] + [rest] * (ncols - 1)

    t = Table(td, colWidths=col_w, repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), HexColor("#1F5C99")),
        ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#AAAAAA")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("LEFTPADDING", (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [HexColor("#FFFFFF"), HexColor("#EEF4FB")]),
    ]))
    story.append(t)
    story.append(Spacer(1, 8))

async def generate_xlsx(data: str, name: str) -> str:
    """Создаёт Excel-файл из markdown-таблицы или простого текста с разделителем |.
    data может содержать несколько таблиц разделённых пустыми строками."""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    import tempfile

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = name[:31] if name else "Лист1"

    HEADER_FILL = PatternFill("solid", fgColor="1F5C99")
    ROW_FILL_ALT = PatternFill("solid", fgColor="EEF4FB")
    HEADER_FONT = Font(bold=True, color="FFFFFF", name="Calibri", size=11)
    BODY_FONT = Font(name="Calibri", size=11)
    thin = Side(style="thin", color="CCCCCC")
    BORDER = Border(left=thin, right=thin, top=thin, bottom=thin)

    current_row = 1
    is_first_row_in_block = True

    for line in data.strip().split("\n"):
        line = line.strip()
        if not line:
            current_row += 1
            is_first_row_in_block = True
            continue
        # Пропускаем строки-разделители |---|---|
        if re.match(r'^[\|\-\s:]+$', line):
            continue
        if "|" in line:
            cells = [c.strip() for c in line.strip("|").split("|")]
            for col_idx, val in enumerate(cells, start=1):
                cell = ws.cell(row=current_row, column=col_idx, value=val)
                cell.border = BORDER
                cell.alignment = Alignment(wrap_text=True, vertical="top")
                if is_first_row_in_block:
                    cell.fill = HEADER_FILL
                    cell.font = HEADER_FONT
                else:
                    fill = ROW_FILL_ALT if current_row % 2 == 0 else PatternFill("solid", fgColor="FFFFFF")
                    cell.fill = fill
                    cell.font = BODY_FONT
            is_first_row_in_block = False
        else:
            cell = ws.cell(row=current_row, column=1, value=line)
            cell.font = BODY_FONT
            is_first_row_in_block = False
        current_row += 1

    # Автоширина колонок
    for col in ws.columns:
        max_len = 0
        col_letter = col[0].column_letter
        for cell in col:
            if cell.value:
                max_len = max(max_len, len(str(cell.value)))
        ws.column_dimensions[col_letter].width = min(max(max_len + 4, 12), 50)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx", prefix="XLS_")
    wb.save(tmp.name)
    return tmp.name
